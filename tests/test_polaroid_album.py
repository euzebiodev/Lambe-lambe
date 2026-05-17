from io import BytesIO

import pytest
from PIL import Image
from docx import Document

import polaroid_album as album


def png_size(data):
    with Image.open(BytesIO(data)) as image:
        return image.size, image.mode


def test_mm_to_pt_conversion():
    assert album._mm_to_pt(25.4) == pytest.approx(72)


def test_crop_central_handles_wide_and_tall_images():
    wide = Image.new("RGB", (400, 200), "red")
    tall = Image.new("RGB", (200, 400), "blue")

    assert album._crop_central(wide, 1).size == (200, 200)
    assert album._crop_central(tall, 1).size == (200, 200)


def test_crop_superior_prefers_top_when_tall():
    image = Image.new("RGB", (100, 300), "white")
    cropped = album._crop_superior(image, 1)

    assert cropped.size == (100, 100)


def test_crop_superior_centers_when_wide():
    image = Image.new("RGB", (300, 100), "white")

    cropped = album._crop_superior(image, 1)

    assert cropped.size == (100, 100)


def test_crop_para_aspecto_uses_detected_face(monkeypatch):
    image = Image.new("RGB", (300, 200), "white")
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: (120, 70, 40, 40))

    cropped = album._crop_para_aspecto(image, 1, usar_rosto=True)

    assert cropped.size == (200, 200)


def test_crop_para_aspecto_can_skip_face_detection(monkeypatch):
    image = Image.new("RGB", (300, 200), "white")

    def fail(_image):
        raise AssertionError("face detection should not be called")

    monkeypatch.setattr(album, "_detectar_rosto", fail)
    assert album._crop_para_aspecto(image, 1, usar_rosto=False).size == (200, 200)


def test_encaixar_no_frame_preserves_frame_size():
    image = Image.new("RGB", (400, 100), "green")

    framed = album._encaixar_no_frame(image, 100, 100)

    assert framed.size == (100, 100)


def test_preparar_foto_tirinha_uses_face_branch(monkeypatch):
    image = Image.new("RGB", (300, 400), "white")
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: (100, 120, 60, 60))

    prepared = album._preparar_foto_tirinha(image, 80, 120)

    assert prepared.size == (80, 120)


def test_generates_polaroid_png_without_face(image_file, monkeypatch):
    path = image_file(size=(200, 100))
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: None)

    data = album.gerar_polaroid(str(path))

    assert png_size(data) == ((928, 1288), "RGBA")


def test_generates_polaroid_png_with_face(image_file, monkeypatch):
    path = image_file(size=(160, 220))
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: (50, 60, 40, 40))

    data = album.gerar_polaroid(str(path))

    assert png_size(data) == ((928, 1288), "RGBA")


def test_generates_photo_booth_preview(image_file, monkeypatch):
    path = image_file(size=(120, 220))
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: None)

    data = album.gerar_preview_foto_maquina(str(path))

    size, mode = png_size(data)
    assert size == (708, 529)
    assert mode == "RGB"


def test_generation_ignores_exif_transpose_errors(image_file, monkeypatch):
    path = image_file(size=(120, 220))
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: None)

    def fail(_image):
        raise RuntimeError("bad exif")

    monkeypatch.setattr(album.ImageOps, "exif_transpose", fail)

    assert png_size(album.gerar_preview_foto_maquina(str(path)))[0] == (708, 529)
    assert png_size(album.gerar_polaroid(str(path)))[0] == (928, 1288)


def test_generates_photo_booth_strip(image_file, monkeypatch):
    paths = [str(image_file(f"foto_{i}.jpg", size=(120, 220))) for i in range(5)]
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: None)

    data = album.gerar_tirinha_foto_maquina(paths)

    assert png_size(data) == ((800, 2400), "RGB")


def test_document_helpers_configure_a4_page():
    document = Document()

    album.configurar_pagina_a4(document, 8)

    section = document.sections[0]
    assert section.page_width.mm == pytest.approx(210, abs=0.1)
    assert section.page_height.mm == pytest.approx(297, abs=0.1)


def test_adicionar_espaco_superior_adds_paragraph_when_content_is_short():
    document = Document()

    album.adicionar_espaco_superior(document, 100)

    assert len(document.paragraphs) == 1


def test_montar_documento_polaroid_creates_docx(image_file, tmp_path, monkeypatch):
    paths = [str(image_file(f"p{i}.jpg", size=(120, 160))) for i in range(4)]
    output = tmp_path / "album.docx"
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: None)

    album.montar_documento(paths, str(output), modo=album.MODO_POLAROID)

    assert output.exists()
    assert len(Document(output).tables) == 1


def test_montar_documento_foto_maquina_creates_docx(image_file, tmp_path, monkeypatch):
    paths = [str(image_file(f"t{i}.jpg", size=(120, 160))) for i in range(5)]
    output = tmp_path / "tirinhas.docx"
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: None)

    album.montar_documento(paths, str(output), modo=album.MODO_FOTO_MAQUINA)

    assert output.exists()
    assert len(Document(output).tables) == 1


def test_montar_documento_foto_maquina_adds_page_breaks(image_file, tmp_path, monkeypatch):
    paths = [str(image_file(f"m{i}.jpg", size=(80, 120))) for i in range(13)]
    output = tmp_path / "multi-pagina.docx"
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: None)

    album.montar_documento(paths, str(output), modo=album.MODO_FOTO_MAQUINA)

    assert output.exists()
    assert len(Document(output).tables) == 2


def test_main_validates_missing_file(monkeypatch, capsys):
    monkeypatch.setattr(
        "sys.argv",
        ["polaroid_album.py", "-o", "saida.docx", "arquivo-inexistente.jpg"],
    )

    with pytest.raises(SystemExit) as exc:
        album.main()

    assert "arquivo nao encontrado" in str(exc.value)


def test_main_generates_output_with_docx_suffix(image_file, tmp_path, monkeypatch):
    path = image_file(size=(80, 120))
    output = tmp_path / "album-sem-extensao"
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: None)
    monkeypatch.setattr("sys.argv", ["polaroid_album.py", "-o", str(output), str(path)])

    album.main()

    assert output.with_suffix(".docx").exists()


def test_main_prints_warning_when_cv2_unavailable(image_file, tmp_path, monkeypatch, capsys):
    path = image_file(size=(80, 120))
    output = tmp_path / "album.docx"
    monkeypatch.setattr(album, "_CV2_OK", False)
    monkeypatch.setattr(album, "_detectar_rosto", lambda _image: None)
    monkeypatch.setattr("sys.argv", ["polaroid_album.py", "-o", str(output), str(path)])

    album.main()

    assert "OpenCV nao instalado" in capsys.readouterr().out
    assert output.exists()
