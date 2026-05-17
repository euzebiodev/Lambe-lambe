"""
polaroid_album.py
-----------------
Gera um documento Word (.docx) com fotos no estilo polaroid,
distribuidas em 3 colunas numa unica pagina A4. A linha de corte
fica desenhada no proprio contorno da moldura (cinza bem claro).

Quando OpenCV esta instalado, o script detecta rostos automaticamente
e centraliza o recorte da polaroid no maior rosto encontrado, com
um leve vies para que o rosto fique no terco superior (em vez de
matematicamente no centro). Se nao houver rosto ou se OpenCV nao
estiver disponivel, cai no comportamento padrao (recorte central).

USO:
    python polaroid_album.py -o saida.docx foto1.jpg foto2.jpg ...
    python polaroid_album.py --modo foto_maquina -o saida.docx foto1.jpg ...

REQUISITOS:
    pip install pillow python-docx
    pip install opencv-python    # para detectar rosto
"""

import argparse
import os
import sys
import tempfile
import warnings
from io import BytesIO

try:
    from PIL import Image, ImageDraw, ImageFilter, ImageOps
except ImportError:
    sys.exit("ERRO: instale a Pillow com:  pip install pillow")

Image.MAX_IMAGE_PIXELS = int(os.environ.get("POLAROID_MAX_IMAGE_PIXELS", "24000000"))
warnings.simplefilter("error", Image.DecompressionBombWarning)

# OpenCV e opcional. Se nao tiver, segue sem deteccao de rosto.
try:
    import cv2
    import numpy as np
    _CV2_OK = True
except ImportError:
    _CV2_OK = False

try:
    from docx import Document
    from docx.shared import Mm, Pt, Emu
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("ERRO: instale o python-docx com:  pip install python-docx")


# --------- Parametros da polaroid ----------
PHOTO_W = 800
PHOTO_H = 1000
BORDER_TOP = 40
BORDER_SIDE = 40
BORDER_BOTTOM = 200
SHADOW_OFFSET = 12
SHADOW_BLUR = 18

# Linha de corte clara no contorno da moldura
LINHA_CORTE_LARGURA = 2
LINHA_CORTE_RGB = (215, 215, 215)

# Onde posicionar verticalmente o rosto dentro do recorte (0=topo, 1=base).
# 0.4 deixa o rosto no terco superior, igual ao "regra dos tercos".
ROSTO_POSICAO_VERTICAL = 0.40

# --------- Parametros do documento ----------
COLS = 3
PAGE_MARGIN_MM = 8
POLAROID_W_MM = 55
POLAROID_IMG_H_MM = POLAROID_W_MM * (
    (PHOTO_H + BORDER_TOP + BORDER_BOTTOM + SHADOW_OFFSET * 4)
    / (PHOTO_W + 2 * BORDER_SIDE + SHADOW_OFFSET * 4)
)
ROW_GAP_MM = 4

# --------- Parametros da tirinha de cabine ----------
MODO_POLAROID = "polaroid"
MODO_FOTO_MAQUINA = "foto_maquina"
FOTOS_POR_TIRINHA = 4
TIRINHAS_POR_PAGINA = 3
TIRINHA_W = 800
TIRINHA_H = 2400
TIRINHA_MARGEM = 46
TIRINHA_ESPACO = 34
TIRINHA_RODAPE = 90
TIRINHA_W_MM = 55
TIRINHA_H_MM = TIRINHA_W_MM * (TIRINHA_H / TIRINHA_W)


def _mm_to_pt(mm):
    return mm * 72 / 25.4


# ---------- Deteccao de rosto ----------
def _detectar_rosto(im_pil):
    """Retorna (x, y, w, h) do maior rosto detectado, ou None."""
    if not _CV2_OK:
        return None
    arr = np.array(im_pil.convert("RGB"))
    gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    cascade_path = cv2.data.haarcascades + "haarcascade_frontalface_default.xml"
    detector = cv2.CascadeClassifier(cascade_path)
    if detector.empty():
        return None
    # minSize proporcional para evitar falsos positivos em fotos grandes
    lado_min = max(60, min(im_pil.size) // 12)
    faces = detector.detectMultiScale(
        gray,
        scaleFactor=1.1,
        minNeighbors=5,
        minSize=(lado_min, lado_min),
    )
    if len(faces) == 0:
        return None
    # maior rosto
    x, y, w, h = max(faces, key=lambda b: int(b[2]) * int(b[3]))
    return int(x), int(y), int(w), int(h)


def _crop_centrado_no_rosto(im, face, target_ratio):
    """Recorta im (PIL) na proporcao target_ratio (W/H) tentando manter
    o rosto centralizado horizontalmente e a 40% do topo verticalmente."""
    img_w, img_h = im.size
    img_ratio = img_w / img_h

    # Maior crop possivel mantendo o aspecto alvo
    if img_ratio > target_ratio:
        new_w = int(img_h * target_ratio)
        new_h = img_h
    else:
        new_w = img_w
        new_h = int(img_w / target_ratio)

    fx, fy, fw, fh = face
    fcx = fx + fw // 2
    fcy = fy + fh // 2

    x0 = fcx - new_w // 2
    y0 = fcy - int(new_h * ROSTO_POSICAO_VERTICAL)

    # Garante que nao sai da imagem
    x0 = max(0, min(x0, img_w - new_w))
    y0 = max(0, min(y0, img_h - new_h))

    return im.crop((x0, y0, x0 + new_w, y0 + new_h))


def _crop_central(im, target_ratio):
    """Recorte central, sem deteccao de rosto."""
    img_w, img_h = im.size
    cur_ratio = img_w / img_h
    if cur_ratio > target_ratio:
        new_w = int(img_h * target_ratio)
        x0 = (img_w - new_w) // 2
        return im.crop((x0, 0, x0 + new_w, img_h))
    new_h = int(img_w / target_ratio)
    y0 = (img_h - new_h) // 2
    return im.crop((0, y0, img_w, y0 + new_h))


def _crop_para_aspecto(im, target_ratio, usar_rosto=True):
    face = _detectar_rosto(im) if usar_rosto else None
    if face is not None:
        return _crop_centrado_no_rosto(im, face, target_ratio)
    return _crop_central(im, target_ratio)


def _crop_superior(im, target_ratio):
    """Recorte com preferencia pela parte de cima, util quando nao ha rosto."""
    img_w, img_h = im.size
    cur_ratio = img_w / img_h
    if cur_ratio > target_ratio:
        new_w = int(img_h * target_ratio)
        x0 = (img_w - new_w) // 2
        return im.crop((x0, 0, x0 + new_w, img_h))

    new_h = int(img_w / target_ratio)
    y0 = 0
    return im.crop((0, y0, img_w, y0 + new_h))


def _encaixar_no_frame(im, frame_w, frame_h):
    foto = ImageOps.contain(im.convert("RGB"), (frame_w, frame_h), Image.LANCZOS)
    fundo = Image.new("RGB", (frame_w, frame_h), "white")
    x = (frame_w - foto.width) // 2
    y = (frame_h - foto.height) // 2
    fundo.paste(foto, (x, y))
    return fundo


def _preparar_foto_tirinha(im, frame_w, frame_h):
    frame_ratio = frame_w / frame_h
    face = _detectar_rosto(im)
    if face is not None:
        return _crop_centrado_no_rosto(im, face, frame_ratio).resize(
            (frame_w, frame_h),
            Image.LANCZOS,
        )

    return _crop_superior(im, frame_ratio).resize((frame_w, frame_h), Image.LANCZOS)


def gerar_preview_foto_maquina(caminho_imagem: str) -> bytes:
    frame_w = TIRINHA_W - 2 * TIRINHA_MARGEM
    frame_h = (
        TIRINHA_H
        - 2 * TIRINHA_MARGEM
        - TIRINHA_RODAPE
        - TIRINHA_ESPACO * (FOTOS_POR_TIRINHA - 1)
    ) // FOTOS_POR_TIRINHA

    im = Image.open(caminho_imagem)
    try:
        im = ImageOps.exif_transpose(im)
    except Exception:
        pass

    foto = _preparar_foto_tirinha(im, frame_w, frame_h).convert("RGB")
    buf = BytesIO()
    foto.save(buf, format="PNG")
    return buf.getvalue()


# ---------- Geracao da polaroid ----------
def gerar_polaroid(caminho_imagem: str) -> bytes:
    im = Image.open(caminho_imagem)
    try:
        im = ImageOps.exif_transpose(im)
    except Exception:
        pass

    target_ratio = PHOTO_W / PHOTO_H

    # Tenta achar rosto na orientacao original (corrigida pelo EXIF)
    face = _detectar_rosto(im)

    if face is None:
        # Sem rosto: roda 90 se horizontal e faz crop central
        if im.width > im.height:
            im = im.rotate(-90, expand=True)
        cropped = _crop_central(im, target_ratio)
    else:
        # Com rosto: nao gira (a foto provavelmente esta na orientacao certa)
        cropped = _crop_centrado_no_rosto(im, face, target_ratio)

    cropped = cropped.resize((PHOTO_W, PHOTO_H), Image.LANCZOS).convert("RGB")

    pola_w = PHOTO_W + 2 * BORDER_SIDE
    pola_h = PHOTO_H + BORDER_TOP + BORDER_BOTTOM
    polaroid = Image.new("RGB", (pola_w, pola_h), "white")
    polaroid.paste(cropped, (BORDER_SIDE, BORDER_TOP))

    # Linha de corte clara no contorno da moldura
    contorno = ImageDraw.Draw(polaroid)
    contorno.rectangle(
        [0, 0, pola_w - 1, pola_h - 1],
        outline=LINHA_CORTE_RGB,
        width=LINHA_CORTE_LARGURA,
    )

    canvas_w = pola_w + SHADOW_OFFSET * 4
    canvas_h = pola_h + SHADOW_OFFSET * 4
    canvas = Image.new("RGBA", (canvas_w, canvas_h), (255, 255, 255, 0))

    sombra = Image.new("RGBA", (canvas_w, canvas_h), (255, 255, 255, 0))
    d = ImageDraw.Draw(sombra)
    off = SHADOW_OFFSET * 2 + SHADOW_OFFSET
    d.rectangle([off, off, off + pola_w, off + pola_h], fill=(0, 0, 0, 100))
    sombra = sombra.filter(ImageFilter.GaussianBlur(SHADOW_BLUR))
    canvas = Image.alpha_composite(canvas, sombra)
    canvas.paste(polaroid, (SHADOW_OFFSET * 2, SHADOW_OFFSET * 2))

    buf = BytesIO()
    canvas.save(buf, format="PNG")
    return buf.getvalue()


def gerar_tirinha_foto_maquina(caminhos_imagens) -> bytes:
    frame_w = TIRINHA_W - 2 * TIRINHA_MARGEM
    frame_h = (
        TIRINHA_H
        - 2 * TIRINHA_MARGEM
        - TIRINHA_RODAPE
        - TIRINHA_ESPACO * (FOTOS_POR_TIRINHA - 1)
    ) // FOTOS_POR_TIRINHA
    frame_ratio = frame_w / frame_h

    tirinha = Image.new("RGB", (TIRINHA_W, TIRINHA_H), "white")
    draw = ImageDraw.Draw(tirinha)

    for idx, caminho in enumerate(caminhos_imagens[:FOTOS_POR_TIRINHA]):
        im = Image.open(caminho)
        try:
            im = ImageOps.exif_transpose(im)
        except Exception:
            pass

        foto = _preparar_foto_tirinha(im, frame_w, frame_h)
        foto = foto.convert("RGB")
        y = TIRINHA_MARGEM + idx * (frame_h + TIRINHA_ESPACO)
        tirinha.paste(foto, (TIRINHA_MARGEM, y))
        draw.rectangle(
            [TIRINHA_MARGEM, y, TIRINHA_MARGEM + frame_w - 1, y + frame_h - 1],
            outline=LINHA_CORTE_RGB,
            width=1,
        )

    draw.rectangle(
        [0, 0, TIRINHA_W - 1, TIRINHA_H - 1],
        outline=LINHA_CORTE_RGB,
        width=LINHA_CORTE_LARGURA,
    )

    buf = BytesIO()
    tirinha.save(buf, format="PNG")
    return buf.getvalue()


# ---------- DOCX ----------
def remover_bordas_tabela(tabela):
    tbl = tabela._tbl
    tblPr = tbl.tblPr
    for existente in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existente)
    tblBorders = OxmlElement("w:tblBorders")
    for borda in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{borda}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def configurar_pagina_a4(documento, margem_mm: float):
    secao = documento.sections[0]
    secao.page_width = Mm(210)
    secao.page_height = Mm(297)
    secao.top_margin = Mm(margem_mm)
    secao.bottom_margin = Mm(margem_mm)
    secao.left_margin = Mm(margem_mm)
    secao.right_margin = Mm(margem_mm)


def adicionar_espaco_superior(documento, altura_conteudo_mm):
    altura_util_mm = 297 - 2 * PAGE_MARGIN_MM
    espaco_mm = max(0, (altura_util_mm - altura_conteudo_mm) / 2)
    if espaco_mm <= 1:
        return

    par = documento.add_paragraph()
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(_mm_to_pt(espaco_mm))
    par.paragraph_format.line_spacing = Pt(1)


def _preparar_documento():
    documento = Document()
    configurar_pagina_a4(documento, PAGE_MARGIN_MM)

    estilo = documento.styles["Normal"]
    estilo.paragraph_format.space_before = Pt(0)
    estilo.paragraph_format.space_after = Pt(0)

    return documento


def _configurar_tabela(tabela, cols):
    tabela.autofit = False
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    remover_bordas_tabela(tabela)

    largura_total = Mm(210 - 2 * PAGE_MARGIN_MM)
    largura_col = Emu(int(largura_total.emu / cols))
    for col in tabela.columns:
        col.width = largura_col

    return largura_col


def montar_documento_polaroid(imagens, saida):
    documento = _preparar_documento()

    num = len(imagens)
    linhas = (num + COLS - 1) // COLS

    tabela = documento.add_table(rows=linhas, cols=COLS)
    largura_col = _configurar_tabela(tabela, COLS)

    arquivos_temp = []
    try:
        for idx, caminho in enumerate(imagens):
            r, c = divmod(idx, COLS)
            cel = tabela.cell(r, c)
            cel.width = largura_col
            cel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            par = cel.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_before = Pt(0)
            par.paragraph_format.space_after = Pt(0)

            print(f"  [{idx+1}/{num}] processando {os.path.basename(caminho)} ...")
            png = gerar_polaroid(caminho)

            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False, prefix="polaroid_")
            tmp.write(png)
            tmp.close()
            arquivos_temp.append(tmp.name)

            run = par.add_run()
            run.add_picture(tmp.name, width=Mm(POLAROID_W_MM))

        documento.save(saida)
    finally:
        for t in arquivos_temp:
            try:
                os.remove(t)
            except OSError:
                pass


def montar_documento_foto_maquina(imagens, saida):
    documento = _preparar_documento()
    grupos = [
        imagens[i:i + FOTOS_POR_TIRINHA]
        for i in range(0, len(imagens), FOTOS_POR_TIRINHA)
    ]
    paginas = (len(grupos) + TIRINHAS_POR_PAGINA - 1) // TIRINHAS_POR_PAGINA

    arquivos_temp = []
    try:
        for pagina in range(paginas):
            if pagina > 0:
                documento.add_page_break()

            grupos_pagina = grupos[
                pagina * TIRINHAS_POR_PAGINA:(pagina + 1) * TIRINHAS_POR_PAGINA
            ]
            adicionar_espaco_superior(documento, TIRINHA_H_MM)
            tabela = documento.add_table(rows=1, cols=TIRINHAS_POR_PAGINA)
            largura_col = _configurar_tabela(tabela, TIRINHAS_POR_PAGINA)

            for idx, grupo in enumerate(grupos_pagina):
                cel = tabela.cell(0, idx)
                cel.width = largura_col
                cel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                par = cel.paragraphs[0]
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                par.paragraph_format.space_before = Pt(0)
                par.paragraph_format.space_after = Pt(0)

                inicio = pagina * TIRINHAS_POR_PAGINA * FOTOS_POR_TIRINHA
                primeira = inicio + idx * FOTOS_POR_TIRINHA + 1
                ultima = primeira + len(grupo) - 1
                print(f"  tirinha {pagina * TIRINHAS_POR_PAGINA + idx + 1}: fotos {primeira}-{ultima} ...")

                png = gerar_tirinha_foto_maquina(grupo)
                tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False, prefix="tirinha_")
                tmp.write(png)
                tmp.close()
                arquivos_temp.append(tmp.name)

                run = par.add_run()
                run.add_picture(tmp.name, width=Mm(TIRINHA_W_MM))

        documento.save(saida)
    finally:
        for t in arquivos_temp:
            try:
                os.remove(t)
            except OSError:
                pass


def montar_documento(imagens, saida, modo=MODO_POLAROID):
    if modo == MODO_FOTO_MAQUINA:
        montar_documento_foto_maquina(imagens, saida)
        return
    montar_documento_polaroid(imagens, saida)


def main():
    parser = argparse.ArgumentParser(
        description="Gera um documento Word com fotos em estilo polaroid ou tirinhas de foto de maquina. Detecta rostos quando OpenCV esta instalado.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Exemplo:\n"
            '  python polaroid_album.py -o album.docx foto1.jpg foto2.jpg foto3.jpg\n'
        ),
    )
    parser.add_argument("-o", "--saida", required=True,
                        help="Caminho do arquivo .docx de saida (incluindo nome).")
    parser.add_argument(
        "--modo",
        choices=(MODO_POLAROID, MODO_FOTO_MAQUINA),
        default=MODO_POLAROID,
        help="Use 'polaroid' para fotos individuais ou 'foto_maquina' para tirinhas verticais com 4 fotos.",
    )
    parser.add_argument("imagens", nargs="+",
                        help="Caminhos das imagens (jpg/png).")
    args = parser.parse_args()

    if not _CV2_OK:
        print("[aviso] OpenCV nao instalado: deteccao de rosto desativada.")
        print("        Instale com:  pip install opencv-python")

    for caminho in args.imagens:
        if not os.path.isfile(caminho):
            sys.exit(f"ERRO: arquivo nao encontrado: {caminho}")

    saida = args.saida
    if not saida.lower().endswith(".docx"):
        saida += ".docx"
    pasta = os.path.dirname(os.path.abspath(saida))
    os.makedirs(pasta, exist_ok=True)

    print(f"Gerando '{saida}' no modo '{args.modo}' com {len(args.imagens)} foto(s)...")
    montar_documento(args.imagens, saida, modo=args.modo)
    print(f"Pronto! Documento salvo em: {saida}")


if __name__ == "__main__":
    main()
